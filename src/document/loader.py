'''
Created on Nov 15, 2013

@author: Spencer Graffe
'''
import os
from PyQt4.QtCore import QThread, pyqtSignal
from PyQt4.QtWebKit import QWebPage
from lxml import html
from docx.docx_document import DocxDocument
from clipboard_document import ClipboardDocument

class DocumentLoadingThread(QThread):
    '''
    This QThread loads in a document asynchronously. It will then tell the GUI
    when it is finished, presenting the document.
    '''
    
    progress = pyqtSignal(int, unicode)
    error = pyqtSignal(object)

    def __init__(self, fileName, isClipboard=False):
        '''
        Constructor
        '''
        QThread.__init__(self)
        
        self._fileName = fileName
        self._isClipboard = isClipboard
        self._success = False
        self._doc = None
        self._canceled = False
    
    def run(self):
        
        self._success = False
        
        if not self._isClipboard:
            # Check which document goes to this extension
            ext = os.path.splitext(os.path.basename(self._fileName))[1].lower()
            if ext == '.docx':
                self._doc = DocxDocument(self._fileName, self._reportProgress, self._isCanceled)
        else:
            self._doc = ClipboardDocument('', self._reportProgress, self._isCanceled)
        
        self._success = not self._isCanceled()
            
    def _reportProgress(self, percent, label):
        self.progress.emit(percent, label)
        
    def _isCanceled(self):
        return self._canceled
    
    def stop(self):
        self._canceled = True
        
    def isSuccess(self):
        return self._success
    
    def getDocument(self):
        if self._success:
            return self._doc
        else:
            return None
        