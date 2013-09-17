'''
Created on Jan 21, 2013

@author: Spencer
'''
import os
import re
import math
import traceback
from lxml import html
from lxml.etree import ParserError, XMLSyntaxError
from HTMLParser import HTMLParser
from PyQt4 import QtGui
from PyQt4.QtCore import Qt, QUrl, QMutex, pyqtSignal, QThread
from PyQt4.QtWebKit import QWebPage, QWebInspector, QWebSettings
from forms.mainwindow_ui import Ui_MainWindow
from gui.bookmarks import BookmarksTreeModel, BookmarkNode
from gui.configuration import Configuration
from gui.npa_webview import NPAWebView
from gui.prepare_speech import PrepareSpeechProgressWidget
from gui.search_settings import SearchSettings
from gui.save_mp3_pages_dialog import SaveMP3PagesDialog
from mathml.tts import MathTTS
from speech.assigner import Assigner, PrepareSpeechThread
from speech.worker import SpeechWorker
import misc
from updater import GetUpdateThread, RunUpdateInstallerThread, SETUP_FILE, SETUP_TEMP_FILE, run_update_installer, is_update_downloaded, save_server_version_to_temp

class MainWindow(QtGui.QMainWindow):
    loc = 0
    len = 0
    jumped = False
    
    # TTS control signals
    startPlayback = pyqtSignal()
    stopPlayback = pyqtSignal()
    setSpeechGenerator = pyqtSignal(object)
    noMoreSpeech = pyqtSignal()
    
    # TTS setting signals
    changeVolume = pyqtSignal(int)
    changeRate = pyqtSignal(int)
    changePauseLength = pyqtSignal(int)
    changeVoice = pyqtSignal(str)
    changeMathDatabase = pyqtSignal(str)
    
    # Program update notification
    notifyProgramUpdate = pyqtSignal()
    programUpdateFinish = pyqtSignal()
    
    # JavaScript mutex
    javascriptMutex = QMutex()
    
    def __init__(self, app, parent=None):
        QtGui.QMainWindow.__init__(self, parent)
        
        self.app = app
        
        self.ui = Ui_MainWindow()
        self.ui.setupUi(self)
        
        # Replace the web view in my designer with my custom web view
        self.ui.webViewLayout.removeWidget(self.ui.webView)
        self.ui.webView.close()
        self.ui.webView = NPAWebView(mainWindow=self, parent=self.ui.centralwidget)
        self.ui.webViewLayout.insertWidget(2, self.ui.webView, stretch=1)
        self.ui.webViewLayout.update()
        
        # Get the search function widgets
        self.searchWidgets = []
        self.searchWidgets.append(self.ui.searchLabel)
        self.searchWidgets.append(self.ui.searchUpButton)
        self.searchWidgets.append(self.ui.searchDownButton)
        self.searchWidgets.append(self.ui.searchTextBox)
        self.searchWidgets.append(self.ui.searchSettingsButton)
        self.searchWidgets.append(self.ui.closeSearchButton)
        self.hideSearch()
        
        # Hide the update button
        self.ui.getUpdateButton.hide()
        
        # Set the search settings dialog so I can make it non-modal
        self.searchSettings = SearchSettings()
        self.searchSettings.setWindowFlags(self.searchSettings.windowFlags() | Qt.WindowStaysOnTopHint)
        
        # Connect all of my signals
        self.connect_signals()
        
        # This is the tree model used to store our bookmarks
        self.bookmarksModel = BookmarksTreeModel(BookmarkNode(None, 'Something'))
        self.ui.bookmarksTreeView.setModel(self.bookmarksModel)
        
        # Hide and other disable development tools if we are in a release
        # environment
        if misc.is_release_environment():
            self.ui.menuMathML.menuAction().setVisible(False)
        else:
            # This is my web inspector for debugging my JavaScript code
            self.webInspector = QWebInspector()
            self.ui.webView.settings().setAttribute(QWebSettings.DeveloperExtrasEnabled, True)
            self.webInspector.setPage(self.ui.webView.page())
        
        # Set the math TTS using the general pattern database
        try:
            dList = misc.pattern_databases()
            self.mathTTS = MathTTS(dList['General'])
        except Exception:
            self.mathTTS = None
            tb = traceback.format_exc()
            print tb            
            message = QtGui.QMessageBox()
            message.setText('The math parser is not working right now. Don\'t read anything math-related for now.')            
        
        # TTS states
        self.resetTTSStates()
        
        # This class is used for assigning the different parts of the content
        # to a specific method of how to speak it
        self.assigner = Assigner()
        
        # Set up my speech driver and start it
        self.speechThread = SpeechWorker()
        
        self.speechThread.onStart.connect(self.onStart)
        self.speechThread.onWord.connect(self.onWord)
        self.speechThread.onEndStream.connect(self.onEndStream)
        self.speechThread.onFinish.connect(self.onSpeechFinished)
        self.speechThread.requestMoreSpeech.connect(self.sendMoreSpeech)
        
        self.startPlayback.connect(self.speechThread.startPlayback)
        self.stopPlayback.connect(self.speechThread.stopPlayback)
        self.setSpeechGenerator.connect(self.speechThread.setSpeechGenerator)
        self.noMoreSpeech.connect(self.speechThread.noMoreSpeech)
        self.changeVolume.connect(self.speechThread.setVolume)
        self.changeRate.connect(self.speechThread.setRate)
        self.changePauseLength.connect(self.speechThread.setPauseLength)
        self.changeVoice.connect(self.speechThread.setVoice)
        self.changeMathDatabase.connect(self.assigner.setMathDatabase)
                
        self.speechThread.start()
                
        # Create the general-purpose progress dialog
        self.progressDialog = QtGui.QProgressDialog('Stuff', 'Cancel', 0, 100, self)
        
        # Load the configuration file
        self.configuration = Configuration()
        self.configuration.loadFromFile(misc.app_data_path('configuration.xml'))
        
        # Check to see if I have a voice. If I don't, grab the first one from
        # the TTS driver
        if len(self.configuration.voice) == 0:
            voiceList = self.speechThread.getVoiceList()
            if len(voiceList) > 0:
                self.configuration.voice = voiceList[0][1]
        self.updateSettings()
        
        # Set the search settings to use the configuration
        self.searchSettings.setConfig(self.configuration)
        
        # Set the zoom of the content view (separate from the bookmarks zoom)
        self.ui.webView.setZoom(self.configuration.zoom_content)
        
        # Used for refreshing the document later to make changes. Also store
        # the data for the document itself.
        self.lastDocumentFilePath = ''
        self.document = None
        self.stopDocumentLoad = False
        
        # Show the tutorial if I user hasn't seen it yet
        if self.configuration.showTutorial:
            self.openTutorial()
            
            # Immediately turn the switch off
            self.configuration.showTutorial = False
            self.updateSettings()
        
        # Run an update check thread
        self.checkUpdateThread = GetUpdateThread(self.showUpdateButton)
        self.checkUpdateThread.start()
        
        # Program updater threads and dialogs
        self.updateInstallProgressDialog = None
        self.updateDoneDialog = None
        self.updator = RunUpdateInstallerThread()
        self.programUpdateFinish.connect(self.finishUpdateDownload)
        
        # Prepare speech widget that pops up every time you press Play
        # Connect all event handlers that may be relevant to update the position
        self.prepareSpeechProgress = PrepareSpeechProgressWidget(self.ui.playButton, self)
        self.ui.splitter.splitterMoved.connect(self.prepareSpeechProgress.updatePos_splitterMoved)
        self.prepareSpeechProgress.hide()
            
    def closeEvent(self, event):
        # Save the current configuration
        self.configuration.zoom_content = self.ui.webView.getZoom()
        self.configuration.saveToFile(misc.app_data_path('configuration.xml'))
        self.speechThread.quit()
        
    def resizeEvent(self, event):
        self.prepareSpeechProgress.updatePos()
        
    def updateGUI(self):
        QtGui.qApp.processEvents()
        
    def connect_signals(self):
        '''
        A method I made to connect all of my signals to the correct functions.
        '''
        # Toolbar buttons
        self.ui.playButton.clicked.connect(self.playSpeech)
        self.ui.pauseButton.clicked.connect(self.stopSpeech)
        self.ui.colorSettingsButton.clicked.connect(self.showColorSettings)
        self.ui.speechSettingsButton.clicked.connect(self.showSpeechSettings)
        self.ui.saveToMP3Button.clicked.connect(self.saveMP3All)
        self.ui.zoomInButton.clicked.connect(self.zoomIn)
        self.ui.zoomOutButton.clicked.connect(self.zoomOut)
        self.ui.zoomResetButton.clicked.connect(self.zoomReset)
        
        # Main Menu
        # File
        self.ui.actionOpen_Docx.triggered.connect(self.showOpenDocxDialog)
        self.ui.actionSave_All_to_MP3.triggered.connect(self.saveMP3All)
        self.ui.actionSave_Selection_to_MP3.triggered.connect(self.saveMP3Selection)
        self.ui.actionQuit.triggered.connect(self.quit)
        
        # Functions
        self.ui.actionPlay.triggered.connect(self.playSpeech)
        self.ui.actionStop.triggered.connect(self.stopSpeech)
        self.ui.actionZoom_In.triggered.connect(self.zoomIn)
        self.ui.actionZoom_Out.triggered.connect(self.zoomOut)
        self.ui.actionReset_Zoom.triggered.connect(self.zoomReset)
        self.ui.actionSearch.triggered.connect(self.toggleSearchBar)
        
        # MP3
        self.ui.actionEntire_Document.triggered.connect(self.saveMP3All)
        self.ui.actionCurrent_Selection.triggered.connect(self.saveMP3Selection)
        self.ui.actionBy_Page.triggered.connect(self.saveMP3ByPage)
        
        # Settings
        self.ui.actionHighlights_Colors_and_Fonts.triggered.connect(self.showColorSettings)
        self.ui.actionSpeech.triggered.connect(self.showSpeechSettings)
        
        # MathML
        self.ui.actionOpen_Pattern_Editor.triggered.connect(self.openPatternEditor)
        self.ui.actionShow_All_MathML.triggered.connect(self.showAllMathML)
        
        # Help
        self.ui.actionTutorial.triggered.connect(self.openTutorial)
        self.ui.actionAbout.triggered.connect(self.openAboutDialog)
        self.ui.actionReport_a_Bug.triggered.connect(self.openReportBugWindow)
        self.ui.actionTake_A_Survey.triggered.connect(self.openSurveyWindow)
        
        # Search bar
        self.ui.searchUpButton.clicked.connect(self.searchBackwards)
        self.ui.searchDownButton.clicked.connect(self.searchForwards)
        self.ui.searchTextBox.returnPressed.connect(self.searchForwards)
        self.ui.searchSettingsButton.clicked.connect(self.openSearchSettings)
        self.ui.closeSearchButton.clicked.connect(self.closeSearchBar)
        
        # Sliders
        self.ui.volumeSlider.valueChanged.connect(self.changeSpeechVolume)
        self.ui.rateSlider.valueChanged.connect(self.changeSpeechRate)
        
        self.ui.actionIncrease_Volume.triggered.connect(self.increaseVolume)
        self.ui.actionDecrease_Volume.triggered.connect(self.decreaseVolume)
        
        self.ui.actionIncrease_Rate.triggered.connect(self.increaseRate)
        self.ui.actionDecrease_Rate.triggered.connect(self.decreaseRate)
        
        # Bookmark and page controls and widgets
        self.ui.bookmarksTreeView.clicked.connect(self.bookmarksTree_clicked)
        self.ui.pagesTreeView.clicked.connect(self.pagesTree_clicked)
        self.ui.bookmarkZoomInButton.clicked.connect(self.bookmarkZoomInButton_clicked)
        self.ui.bookmarkZoomOutButton.clicked.connect(self.bookmarkZoomOutButton_clicked)
        #self.ui.expandBookmarksButton.clicked.connect(self.expandBookmarksButton_clicked)
        #self.ui.collapseBookmarksButton.clicked.connect(self.collapseBookmarksButton_clicked)
        
        # Update button
        self.ui.getUpdateButton.clicked.connect(self.runUpdate)
        
    def updateSettings(self):
        
        # Update speech thread with my stuff
        self.changeVolume.emit(self.configuration.volume)
        self.changeRate.emit(self.configuration.rate)
        self.changePauseLength.emit(self.configuration.pause_length)
        self.changeVoice.emit(self.configuration.voice)
        
        # Update the math database to use
        try:
            self.changeMathDatabase.emit(misc.pattern_databases()[self.configuration.math_database])
        except KeyError:
            self.changeMathDatabase.emit(misc.pattern_databases()['General'])
        
        try:
            if self.document is not None:
                self.assigner.prepare(self.document.getMainPage())
        except AttributeError:
            pass
        
        # Update main window sliders to match
        self.ui.rateSlider.setValue(self.configuration.rate)
        self.ui.volumeSlider.setValue(int(self.configuration.volume))
        
        # Zoom settings
        currentFont = self.ui.bookmarksTreeView.font()
        currentFont.setPointSize(self.configuration.zoom_navigation_ptsize)
        self.ui.bookmarksTreeView.setFont(currentFont)
        self.ui.pagesTreeView.setFont(currentFont)
        
        # Finally, save it all to file
        self.configuration.saveToFile(misc.app_data_path('configuration.xml'))
        
    def hideSearch(self):
        for w in self.searchWidgets:
            w.hide()
        #self.ui.webView.page().mainFrame().evaluateJavaScript(misc.js_command('ClearAllHighlights', []))
    
    def showSearch(self):
        for w in self.searchWidgets:
            w.show()
        self.ui.searchTextBox.setFocus()
        self.ui.searchTextBox.selectAll()
        
    def openSearchSettings(self):
        self.searchSettings.show()
            
    def playSpeech(self):
        
        # Stop whatever the speech thread may be saying
        self.stopPlayback.emit()
        
        # Reset the TTS states
        self.resetTTSStates()
        self.setSettingsEnableState(False)
        
        # Set the beginning of the streamer
        contents = unicode(self.ui.webView.page().mainFrame().evaluateJavaScript(misc.js_command('SetStreamBeginning', [])).toString())
        
        # Convert the HTML to DOM
        root = None
        try:
            root = html.fromstring(contents)
        except XMLSyntaxError:
            root = html.Element('p')
        
        # Set the speech generator and start playback
        self.setSpeechGenerator.emit(self.assigner.generateSpeech(root, self.configuration))
        self.startPlayback.emit()
    
    def stopSpeech(self):
        self.stopPlayback.emit()
        self.setSettingsEnableState(True)
        self.ui.webView.setFocus()
        
    def toggleSpeech(self):
        '''
        Toggles the state of speech playback
        '''
        if self.speechThread.isPlaying():
            self.stopSpeech()
        else:
            self.playSpeech()
        
    def resetTTSStates(self):
        self.ttsStates = {'lastElement' : ['', -1, '', -1], 
                          'hasWorded' : False}
    
    def onStart(self, offset, length, label, stream, word):
        #print 'window: OnStart;', offset, length, label, stream, word
        self.ui.webView.page().mainFrame().evaluateJavaScript(misc.js_command('StartHighlighting', []))
        
    def onWord(self, offset, length, label, stream, word, isFirst):
        #print 'window: OnWord; offset', offset, ',length', length, ', label', label, ', stream', stream, ', word', [word], ', isFirst', isFirst
        self.hasWorded = True
        lastElement = self.ttsStates['lastElement']
        
        if label == 'text':
            self.ui.webView.page().mainFrame().evaluateJavaScript(misc.js_command('HighlightNextWord', [self.configuration.highlight_line_enable, unicode(word), offset, length]))
            
        elif label == 'math':
            if label != lastElement[2] or stream != lastElement[3] and (lastElement[3] >= 0):
                self.ui.webView.page().mainFrame().evaluateJavaScript(misc.js_command('HighlightNextMath', [self.configuration.highlight_line_enable]))
                    
        elif label == 'image':
            if label != lastElement[2] or stream != lastElement[3] and (lastElement[3] >= 0):
                self.ui.webView.page().mainFrame().evaluateJavaScript(misc.js_command('HighlightNextImage', [self.configuration.highlight_line_enable]))
        else:
            print 'ERROR: I don\'t know what this label refers to for highlighting:', label
        
        self.ttsStates['lastElement'] = [offset, length, label, stream]
        
    def onEndStream(self, stream, label):
        #print 'window: OnEndStream'
        pass
        
    def onSpeechFinished(self):
        #print 'window: OnSpeechFinished'
        
        # Tell JavaScript to not highlight further
        self.ui.webView.page().mainFrame().evaluateJavaScript(misc.js_command('StopHighlighting', []))
        
        self.setSettingsEnableState(True)
        self.resetTTSStates()
        self.ui.webView.setFocus()
        
    def sendMoreSpeech(self):
        '''
        This is response from requestMoreSpeech from the SpeechWorker thread.
        It will either get more speech for the TTS and send it, or tell it that
        no more speech is available.
        '''
        hasMoreSpeech = self.ui.webView.page().mainFrame().evaluateJavaScript(misc.js_command('HasMoreElements', [])).toBool()
        if hasMoreSpeech:
            nextContent = unicode(self.ui.webView.page().mainFrame().evaluateJavaScript(misc.js_command('StreamNextElement', [])).toString())
            
            # Create the HTML DOM from content
            elem = None
            if len(nextContent) > 0:
                try:
                    isMatch = re.search(r'<.+>', nextContent)
                    if isMatch is None:
                        elem = html.Element('p')
                        h = HTMLParser()
                        elem.text = h.unescape(nextContent)
                    else:
                        elem = html.fromstring(nextContent)
                except ParserError as e:
                    elem = html.Element('p')
            else:
                elem = html.Element('p')
            
            # Create and send the speech generator
            self.setSpeechGenerator.emit(self.assigner.generateSpeech(elem, self.configuration))
            
        else:
            self.noMoreSpeech.emit()
            
    def changeSpeechRate(self, value):
        self.configuration.rate = value
        self.changeRate.emit(self.configuration.rate)
        
    def increaseRate(self):
        self.configuration.rate += 10
        if self.configuration.rate > 100:
            self.configuration.rate = 100
        self.ui.rateSlider.setValue(self.configuration.rate)
    
    def decreaseRate(self):
        self.configuration.rate -= 10
        if self.configuration.rate < 0:
            self.configuration.rate = 0
        self.ui.rateSlider.setValue(self.configuration.rate)

    def changeSpeechVolume(self, value):
        self.configuration.volume = value
        self.changeVolume.emit(self.configuration.volume)
    
    def increaseVolume(self):
        self.configuration.volume += 10
        if self.configuration.volume > 100:
            self.configuration.volume = 100
        self.ui.volumeSlider.setValue(self.configuration.volume)
    
    def decreaseVolume(self):
        self.configuration.volume -= 10
        if self.configuration < 0:
            self.configuration.volume = 0
        self.ui.volumeSlider.setValue(self.configuration.volume)
        
    def showColorSettings(self):
        from gui.color_settings import ColorSettings
        dialog = ColorSettings(self)
        result = dialog.exec_()
        self.configuration.loadFromFile(misc.app_data_path('configuration.xml'))
        if result == ColorSettings.RESULT_NEED_REFRESH:
            self.refreshDocument()
        self.updateSettings()
        
    def showSpeechSettings(self):

        # Disconnect my highlighter signals
        self.speechThread.onStart.disconnect(self.onStart)
        self.speechThread.onWord.disconnect(self.onWord)
        self.speechThread.onEndStream.disconnect(self.onEndStream)
        self.speechThread.onFinish.disconnect(self.onSpeechFinished)
        self.speechThread.requestMoreSpeech.disconnect(self.sendMoreSpeech)

        # Show speech settings dialog
        from gui.speech_settings import SpeechSettings
        dialog = SpeechSettings(self)
        self.speechThread.requestMoreSpeech.connect(dialog.requestMoreSpeech)
        dialog.exec_()
        self.stopPlayback.emit()
        self.speechThread.requestMoreSpeech.disconnect(dialog.requestMoreSpeech)
        self.configuration.loadFromFile(misc.app_data_path('configuration.xml')) 
        self.updateSettings()

        # Reconnect my highlighter signals
        self.speechThread.onStart.connect(self.onStart)
        self.speechThread.onWord.connect(self.onWord)
        self.speechThread.onEndStream.connect(self.onEndStream)
        self.speechThread.onFinish.connect(self.onSpeechFinished)
        self.speechThread.requestMoreSpeech.connect(self.sendMoreSpeech)
        
    def saveMP3All(self):
        self.ui.webView.page().mainFrame().evaluateJavaScript(misc.js_command('ClearAllHighlights', []))
        self.ui.webView.selectAll()
        self.saveMP3Selection()
        self.ui.webView.page().mainFrame().evaluateJavaScript(misc.js_command('SetHighlightToBeginning', []))
        
    def saveMP3Selection(self):
        # Generate a filename that is basically the original file but with the
        # .mp3 extension at the end
        defaultFileName = os.path.splitext(str(self.lastDocumentFilePath))[0] + '.mp3'
        fileName = QtGui.QFileDialog.getSaveFileName(self, 'Save MP3...', defaultFileName, '(*.mp3)')[0]
        print 'Saving to...', fileName
        
        if len(fileName) > 0:
            # Show a progress dialog
            self.progressDialog.setWindowModality(Qt.WindowModal)
            self.progressDialog.setWindowTitle('Saving to MP3...')
            self.progressDialog.setLabelText('Generating speech...')
            self.progressDialog.setProgress(0)
            self.progressDialog.show()
            QtGui.qApp.processEvents()
            
            # Get my speech output list
            selectedHTML = unicode(self.ui.webView.page().mainFrame().evaluateJavaScript(misc.js_command('GetSelectionHTML', [])).toString())
            speechGenerator = self.assigner.generateSpeech(html.fromstring(selectedHTML), self.configuration)
            
            # Get the progress of the thing from the speech thread
            def myOnProgress(percent):
                self.progressDialog.setProgress(percent)
                QtGui.qApp.processEvents()
                
            def myOnProgressLabel(newLabel):
                self.progressDialog.setLabelText(newLabel)
                QtGui.qApp.processEvents()
                
            self.speechThread.onProgress.connect(myOnProgress)
            self.speechThread.onProgressLabel.connect(myOnProgressLabel)
            self.progressDialog.canceled.connect(self.speechThread.stopMP3)
            self.speechThread.saveToMP3(fileName, speechGenerator)
            
            # Just hide it so that we can use it later
            self.progressDialog.hide()
            
            # Show a message box saying the file was successfully saved
            if not self.speechThread.mp3Interrupted():
                messageBox = QtGui.QMessageBox()
                messageText = 'Success!\nYour MP3 was saved as:\n' + fileName
                messageBox.setText(messageText)
                messageBox.setStandardButtons(QtGui.QMessageBox.Ok)
                messageBox.setDefaultButton(QtGui.QMessageBox.Ok)
                messageBox.setWindowTitle('MP3 Saved Successfully')
                messageBox.setIcon(QtGui.QMessageBox.Information)
                messageBox.exec_()
                
                misc.open_file_browser_to_location(fileName)
        
        self.ui.webView.page().mainFrame().evaluateJavaScript(misc.js_command('ClearUserSelection', []))
    
    def saveMP3ByPage(self):
        
        if self.document is not None:
            # Open a dialog to select the folder
            folder = unicode(QtGui.QFileDialog.getExistingDirectory(self, 
                                                            'Select Folder to Save Pages',
                                                            os.path.join(os.path.expanduser('~'), 'Desktop')))
            print 'Saving pages to:', folder
            
            if len(folder) > 0:
                
                # Show the dialog that facilitates the per-page export
                success = True
                dialog = SaveMP3PagesDialog(self)
                dialog.show()
                
                myHTML = html.fromstring(self.document.getMainPage())
                body = myHTML.xpath('//body')[0]
                
                # Split up the body into a list of elements, split by page.
                pages = []
                currentPage = [[], 'Front']
                largestNumberLength = 1
                for e in body:
                    if e.get('class') == 'pageNumber':
                        
                        if len(e.text_content()) > 0:
                            
                            # Figure out the number of decimal places of the page number, if possible
                            try:
                                num = math.ceil(math.log10(int(e.text_content())))
                                if num > largestNumberLength:
                                    largestNumberLength = int(num)
                            except ValueError:
                                pass
                                
                        if len(currentPage[0]) > 0:
                            pages.append(currentPage)
                        currentPage = [[], e.text_content()]
                    currentPage[0].append(e)
                
                if len(currentPage[0]) > 0:
                    pages.append(currentPage)
                
                # Save all of the pages!
                for i in range(len(pages)):
                    page = pages[i]
                    
                    # Tell the dialog what page we are on
                    dialog.setPageLabel('Page ' + page[1] + '...')
                    dialog.setPageProgress(int((i+1.0) / len(pages) * 100.0))
                    QtGui.qApp.processEvents()
                    
                    # Prepare the speech of the page
                    dialog.setStatusLabel('Preparing...')
                    dialog.setStatusProgress(0)
                    QtGui.qApp.processEvents()
                    myHTML = ''
                    for e in page[0]:
                        myHTML += html.tostring(e)
                        
                    speechGenerator = self.assigner.generateSpeech(html.fromstring(myHTML), self.configuration)
                
                    # Generate a file name for saving the page number
                    fileName = ''
                    if misc.is_number(page[1]):
                        fileName = 'Pg ' + page[1].zfill(largestNumberLength) + ' ' + os.path.splitext(os.path.basename(self.lastDocumentFilePath))[0]
                    else:
                        fileName = 'Pg ' + page[1] + ' ' + os.path.splitext(os.path.basename(self.lastDocumentFilePath))[0]
                    fileName += '.mp3'
                    fileName = os.path.join(folder, fileName)
                
                    # Get the progress of the thing from the speech thread
                    def myOnProgress(percent):
                        dialog.setStatusProgress(percent)
                        QtGui.qApp.processEvents()
                        
                    def myOnProgressLabel(newLabel):
                        dialog.setStatusLabel(newLabel)
                        QtGui.qApp.processEvents()
                    
                    self.speechThread.onProgress.connect(myOnProgress)
                    self.speechThread.onProgressLabel.connect(myOnProgressLabel)
                    dialog.canceled.connect(self.speechThread.stopMP3)
                    
                    self.speechThread.saveToMP3(fileName, speechGenerator)
                    
                    if self.speechThread.mp3Interrupted():
                        success = False
                        break
    
                dialog.close()
                
                # Show the success message, if I was indeed successful
                if success:
                    messageBox = QtGui.QMessageBox()
                    messageText = 'Success!\nYour MP3 files were saved in:\n' + folder
                    messageBox.setText(messageText)
                    messageBox.setStandardButtons(QtGui.QMessageBox.Ok)
                    messageBox.setDefaultButton(QtGui.QMessageBox.Ok)
                    messageBox.setWindowTitle('MP3 Files Saved Successfully')
                    messageBox.setIcon(QtGui.QMessageBox.Information)
                    messageBox.exec_()
                    
                    misc.open_file_browser_to_location(folder)
        
    def zoomIn(self):
        self.ui.webView.zoomIn()
    
    def zoomOut(self):
        self.ui.webView.zoomOut()
        
    def zoomReset(self):
        self.ui.webView.zoomReset()
            
#     def openHTML(self):
#         fileName = QtGui.QFileDialog.getOpenFileName(self, 'Open HTML...','./tests','(*.html)')
#         f = open(fileName, 'r')
#         content = f.read()
#         f.close()
#         self.assigner.prepare(content)
#         baseUrl = QUrl.fromLocalFile(fileName)
#         self.ui.webView.setHtml(content, baseUrl)


    def openDocx(self, filePath):
        '''
        Opens a .docx file. It will spawn a progress bar and run the task in the
        background to prevent GUI from freezing up.
        '''
        filePath = str(filePath)
        if len(filePath) > 0:
            
            from docx.thread import DocxImporterThread
            from gui.document_load_progress import DocumentLoadProgressDialog
             
            # Create my .docx importer thread
            self.docxImporterThread = DocxImporterThread(filePath)
            self.docxImporterThread.reportProgress.connect(self.reportProgressOpenDocx)
            self.docxImporterThread.reportError.connect(self.reportErrorOpenDocx)
            self.docxImporterThread.finished.connect(self.finishOpenDocx)
            
            # Show a progress dialog
            self.progressDialog = DocumentLoadProgressDialog()
            self.progressDialog.setLabelText('Reading ' + os.path.basename(str(filePath)) + '...')
            self.progressDialog.setProgress(0)
            self.progressDialog.canceled.connect(self.docxImporterThread.stop)
            self.progressDialog.show()
           
            self.docxImporterThread.start()
        
    def reportProgressOpenDocx(self, percent):
        self.progressDialog.setProgress(percent - 1)
        
    def reportTextOpenDocx(self, text):
        self.progressDialog.setLabelText(text)
        
    def reportErrorOpenDocx(self, exception, tb):
        from mathtype.parser import MathTypeParseError
        from gui.bug_reporter import BugReporter
        
        if isinstance(exception, MathTypeParseError):
            out = misc.prepare_bug_report(tb, self.configuration, detailMessage=exception.message)
            dialog = BugReporter(out)
            dialog.exec_()
        else:
            out = misc.prepare_bug_report(tb, self.configuration)
            dialog = BugReporter(out)
            dialog.exec_()
            
        if self.progressDialog is not None:
            self.progressDialog.enableCancel()
            self.progressDialog.close()
        
    def finishOpenDocx(self):
        
        print 'Finishing up...'
        
        if self.docxImporterThread.isSuccessful():
            
            # Disable the cancel in the progress thing, since we can't cancel
            # here anyways
            self.progressDialog.disableCancel()
            
            url = misc.temp_path('import')
            baseUrl = QUrl.fromLocalFile(url)
            
            # Clear the cache in the web view
            QWebSettings.clearIconDatabase()
            QWebSettings.clearMemoryCaches()
            
            # Set the content views and prepare assigner
            self.progressDialog.setLabelText('Loading content into view...')
            
            docxHtml = self.docxImporterThread.getHTML()
            self.assigner.prepare(docxHtml)
            self.ui.webView.loadProgress.connect(self.progressDialog.setProgress)
            
            # Use the web view to figure out when the view is done loading
            self.pageLoaded = False
            def setLoadedFinished():
                self.pageLoaded = True
            self.ui.webView.loadFinished.connect(setLoadedFinished)
            self.ui.webView.setHtml(docxHtml, baseUrl)
            
            # Get and set the bookmarks
            self.bookmarksModel = BookmarksTreeModel(self.docxImporterThread.getHeadings())
            self.ui.bookmarksTreeView.setModel(self.bookmarksModel)
            self.ui.bookmarksTreeView.expandAll()
                    
            # Get and set the pages
            from gui.pages import PagesTreeModel
            
            self.pagesModel = PagesTreeModel(self.docxImporterThread.getPages())
            self.ui.pagesTreeView.setModel(self.pagesModel)
            self.ui.pagesTreeView.expandAll()
            
            self.document = self.docxImporterThread.getDocument()
            self.lastDocumentFilePath = self.docxImporterThread.getFilePath()
            
            # Wait until page is done loading
            while not self.pageLoaded:
                QtGui.qApp.processEvents()
            self.progressDialog.enableCancel()
             
            # Wait until MathJax is done typesetting
            self.progressDialog.setLabelText('Typesetting math equations...')
            self.mathjax_loaded = False
             
            # Allow the user to cancel this. Sometimes MathJax freaks out and
            # the user should say no to it
            self.progressDialog.enableCancel()
            def myCancelHandler():
                self.mathjax_loaded = True
            self.progressDialog.canceled.connect(myCancelHandler)
             
            while not self.mathjax_loaded:
                QtGui.qApp.processEvents()
                progress = int(self.ui.webView.page().mainFrame().evaluateJavaScript(misc.js_command('GetMathTypesetProgress', [])).toInt()[0])
                self.progressDialog.setProgress(progress)
                self.mathjax_loaded = self.ui.webView.page().mainFrame().evaluateJavaScript(misc.js_command('IsMathTypeset', []))

        self.progressDialog.close()
        self.stopDocumentLoad = False
        self.ui.webView.setFocus()
        QtGui.qApp.processEvents()
        
    def showOpenDocxDialog(self):
        filePath = QtGui.QFileDialog.getOpenFileName(self, 'Open Docx...',os.path.join(os.path.expanduser('~'), 'Documents'),'(*.docx)')
        self.openDocx(filePath)
        
    def showDocNotSupported(self):
        result = QtGui.QMessageBox.information(self, 'CAR does not support 1997-2003 Word document', 
                                               'Sorry, this is a 1997-2003 Word file than we don\'t support. Open it in Microsoft Word 2007 or later and save it as "Word Document."' 
                                               )
            
    def openTutorial(self):
        self.openDocx(misc.program_path('Tutorial.docx'))
            
    def openAboutDialog(self):
        from gui.about import AboutDialog
        dialog = AboutDialog()
        dialog.exec_()
        
    def openReportBugWindow(self):
        import webbrowser
        webbrowser.open_new(misc.REPORT_BUG_URL)
        
    def openSurveyWindow(self):
        import webbrowser
        webbrowser.open_new(misc.SURVEY_URL)
        
    def toggleSearchBar(self):
        if self.searchWidgets[0].isHidden():
            self.showSearch()
        else:
            self.hideSearch()
            
    def searchBackwards(self):
        text = unicode(self.ui.searchTextBox.text())
        args = [text, False, self.configuration.search_whole_word, self.configuration.search_match_case]
        result = self.ui.webView.page().mainFrame().evaluateJavaScript(misc.js_command('SearchForText', args))
        if not result:
            self.ui.webView.page().mainFrame().evaluateJavaScript(misc.js_command('ClearAllHighlights', []))
            message = QtGui.QMessageBox()
            message.setText('No other occurrences of "' + text + '" in document.')
            message.exec_()
    
    def searchForwards(self):
        text = unicode(self.ui.searchTextBox.text())
        args = [text, True, self.configuration.search_whole_word, self.configuration.search_match_case]
        result = self.ui.webView.page().mainFrame().evaluateJavaScript(misc.js_command('SearchForText', args))
        if not result:
            self.ui.webView.page().mainFrame().evaluateJavaScript(misc.js_command('ClearAllHighlights', []))
            message = QtGui.QMessageBox()
            message.setText('No other occurrences of "' + text + '" in document.')
            message.exec_()
            
    def closeSearchBar(self):
        self.hideSearch()
            
    def quit(self):
        self.close()
        
    def openPatternEditor(self):
        
        from mathml.pattern_editor.gui.patterneditorwindow import PatternEditorWindow

        patternFilePath = self.configuration.math_database
        
        self.patternWindow = PatternEditorWindow(patternFilePath, '')
        self.patternWindow.changedPattern.connect(self.onChangedPatternEditor)
        self.patternWindow.show()
        
    def onChangedPatternEditor(self, databaseFileName):
        self.mathTTS.setPatternDatabase(databaseFileName)
        
    def showAllMathML(self):
        from gui.mathmlcodes_dialog import MathMLCodesDialog
        self.mathmlDialog = MathMLCodesDialog(self.assigner._maths)
        self.mathmlDialog.show()
        
    def bookmarksTree_clicked(self, index):
        node = index.internalPointer()
        self.ui.webView.page().mainFrame().evaluateJavaScript(misc.js_command('GotoPageAnchor', [node.anchorId]))
        
    def pagesTree_clicked(self, index):
        node = index.internalPointer()
        self.ui.webView.page().mainFrame().evaluateJavaScript(misc.js_command('GotoPageAnchor', [node.anchorId]))
        
    def bookmarkZoomInButton_clicked(self):
        # Get the current font
        currentFont = self.ui.bookmarksTreeView.font()
        
        # Make it a litter bigger
        self.configuration.zoom_navigation_ptsize += 2
        currentFont.setPointSize(self.configuration.zoom_navigation_ptsize)
        
        # Set the font
        self.ui.bookmarksTreeView.setFont(currentFont)
        self.ui.pagesTreeView.setFont(currentFont)
        
    def bookmarkZoomOutButton_clicked(self):
        # Get the current font
        currentFont = self.ui.bookmarksTreeView.font()
        
        # Make it a litter smaller
        self.configuration.zoom_navigation_ptsize -= 2
        if self.configuration.zoom_navigation_ptsize < 4:
            self.configuration.zoom_navigation_ptsize = 4
        currentFont.setPointSize(self.configuration.zoom_navigation_ptsize)
        
        # Set the font
        self.ui.bookmarksTreeView.setFont(currentFont)
        self.ui.pagesTreeView.setFont(currentFont)
        
    def expandBookmarksButton_clicked(self):
        self.ui.bookmarksTreeView.expandAll()
        
    def collapseBookmarksButton_clicked(self):
        self.ui.bookmarksTreeView.collapseAll()
        
    def refreshDocument(self):
        
        if self.document != None:
            url = misc.temp_path('import')
            baseUrl = QUrl.fromLocalFile(url)
            
            docxHtml = self.document.getMainPage()
                        
            # Clear the cache in the web view
            QWebSettings.clearIconDatabase()
            QWebSettings.clearMemoryCaches()
                        
            # Set the content views and prepare assigner
            self.assigner.prepare(docxHtml)
            self.ui.webView.setHtml(docxHtml, baseUrl)
        
#         if len(self.lastDocumentFilePath) > 0:
#             self.openDocx(self.lastDocumentFilePath)

    def showUpdateButton(self):
        '''
        Shows the Update button if there is an update available. This should
        only be done by the thread that checks for it.
        '''
        self.ui.getUpdateButton.show()
        
    def runUpdate(self):
        '''
        Runs the setup file to update this program, when one exists.
        '''
        # Depending on whether the update has already downloaded or not, make
        # the prompt tell the user the appropriate action (download it or
        # install it)
        from gui.update_prompt import UpdatePromptDialog
        
        try:
            if is_update_downloaded():
                question = UpdatePromptDialog(self)
                question.setText('Update has already downloaded. Want to install it?')
                result = question.exec_()
                
                if result == QtGui.QMessageBox.Yes:
                    # Run the installer and close this window down
                    run_update_installer()
                    self.close()
                                
            else:
                question = UpdatePromptDialog(self)
                result = question.exec_()
            
                if result == QtGui.QMessageBox.Yes:
                    
                    from gui.download_progress import DownloadProgressWidget
                    
                    # Create the widget for the download right below the content 
                    # view
                    self.updateDownloadProgress = DownloadProgressWidget(self)
                    self.updateDownloadProgress.setUrl(SETUP_FILE)
                    self.updateDownloadProgress.setDestination(SETUP_TEMP_FILE)
                
                    self.updateDownloadProgress.downloadFinished.connect(self.finishUpdateDownload)
                    self.ui.webViewLayout.addWidget(self.updateDownloadProgress, stretch=0)
                
                    self.updateDownloadProgress.startDownload()
                    
        except Exception as e:
            # Generate a bug report for it
            import traceback
            from misc import prepare_bug_report
            from gui.bug_reporter import BugReporter
            out = prepare_bug_report(traceback.format_exc(), self.configuration)
            dialog = BugReporter(out)
            dialog.exec_()
            
    def finishUpdateDownload(self, success):
        '''
        Does the cleanup logic for downloading an update.
        '''
        self.ui.webViewLayout.removeWidget(self.updateDownloadProgress)
        self.updateDownloadProgress.hide()
        if success:
            
            # Save the version from over there to here
            versionInfo = save_server_version_to_temp()
            
            # Prompt the user if they want to install it
            if len(versionInfo) > 0:
                from gui.update_prompt import UpdatePromptDialog
                question = UpdatePromptDialog(self)
                question.setText('Downloaded update! Want to install it?')
                result = question.exec_()
                if result == QtGui.QMessageBox.Yes:
                    
                    # Run the installer
                    run_update_installer()
                    self.app.exit(0)
            
    def setSettingsEnableState(self, isEnable):
        '''
        Disables or enables the other widgets that shouldn't be active during
        TTS playback.
        '''
        if not isEnable:
            self.ui.colorSettingsButton.setEnabled(False)
            self.ui.speechSettingsButton.setEnabled(False)
            self.ui.saveToMP3Button.setEnabled(False)
            self.ui.playButton.setEnabled(False)
            
            # Disable slider bars if TTS is not interactive
            if not self.speechThread.areSettingsInteractive():
                self.ui.rateSlider.setEnabled(False)
                self.ui.volumeSlider.setEnabled(False)
            
            # Actions
            self.ui.actionPlay.setEnabled(False)
            self.ui.actionHighlights_Colors_and_Fonts.setEnabled(False)
            self.ui.actionSpeech.setEnabled(False)
            self.ui.actionSave_All_to_MP3.setEnabled(False)
            self.ui.actionSave_Selection_to_MP3.setEnabled(False)
            self.ui.actionOpen_Docx.setEnabled(False)
            self.ui.actionTutorial.setEnabled(False)
            self.ui.actionSearch.setEnabled(False)
            
            # Search bar
            for w in self.searchWidgets:
                w.setEnabled(False)
                
            # Cursor
            self.ui.webView.setKeyboardNavEnabled(False)
            
        else:
            self.ui.rateSlider.setEnabled(True)
            self.ui.volumeSlider.setEnabled(True)
            self.ui.colorSettingsButton.setEnabled(True)
            self.ui.speechSettingsButton.setEnabled(True)
            self.ui.saveToMP3Button.setEnabled(True)
            self.ui.playButton.setEnabled(True)
            
            # Actions
            self.ui.actionPlay.setEnabled(True)
            self.ui.actionHighlights_Colors_and_Fonts.setEnabled(True)
            self.ui.actionSpeech.setEnabled(True)
            self.ui.actionSave_All_to_MP3.setEnabled(True)
            self.ui.actionSave_Selection_to_MP3.setEnabled(True)
            self.ui.actionOpen_Docx.setEnabled(True)
            self.ui.actionTutorial.setEnabled(True)
            self.ui.actionSearch.setEnabled(True)
            
            # Search bar
            for w in self.searchWidgets:
                w.setEnabled(True)
            
            # Cursor
            self.ui.webView.setKeyboardNavEnabled(True)
